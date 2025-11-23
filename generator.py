import json
import re
import docx
from google import genai
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ======================
#  CONFIGURAÇÃO GERAL
# ======================

API_KEY = "API KEY DO GEMINI AQUI"  # <<< COLOQUE SUA CHAVE AQUI
client = genai.Client(api_key=API_KEY)

REQUISITOS_DOCX = "requisitos.docx"
SAIDA_JSON = "saida.json"
TESTES_JSON = "testes.json"
WORD_OUTPUT = "cenarios_de_testes.docx"

print("🚀 Iniciando gerador de cenários de teste...")

# =============================================
#  1. EXTRAÇÃO DO .DOCX → JSON DE REQUISITOS
# =============================================

def extrair_requisitos_docx(caminho=REQUISITOS_DOCX):
    print("📄 Extraindo requisitos do DOCX...")

    doc = docx.Document(caminho)

    sections = []
    sec_atual = {"title": "", "requirements": []}

    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue

        # Títulos são negrito
        if para.runs and any(r.bold for r in para.runs):
            if sec_atual["title"]:
                sections.append(sec_atual)
            sec_atual = {"title": texto, "requirements": []}
        else:
            sec_atual["requirements"].append({"text": texto})

    if sec_atual["title"]:
        sections.append(sec_atual)

    # salvar JSON
    with open(SAIDA_JSON, "w", encoding="utf-8") as f:
        json.dump({"sections": sections}, f, ensure_ascii=False, indent=2)

    print("✅ Requisitos extraídos e salvos em saida.json")

    return sections

# =================================================================================================
#  2. PROMP PARA GERAÇÃO DOS CENÁRIOS (PODE SER AJUSTADO E MELHORADO CONFORME A SUA NECESSIDADE)
# =================================================================================================

QA_PROMPT = """
Você é um Engenheiro de QA Sênior com 15+ anos de experiência.

## 🎯 OBJETIVO
Gerar suíte COMPLETA de testes baseada nos requisitos fornecidos.

## 📋 METODOLOGIA

### 1. ANÁLISE PRELIMINAR
Identifique:
- Entidades principais e relacionamentos
- Regras de negócio críticas
- Pontos de integração
- Riscos técnicos e de negócio

### 2. TÉCNICAS APLICADAS
- Particionamento de Equivalência
- Análise de Valor Limite
- Tabela de Decisão
- Testes de Estado
- Testes Exploratórios

### 3. COBERTURA OBRIGATÓRIA
✅ Happy path
✅ Validações (formato, tipo, tamanho, regex)
✅ Permissões e autenticação
✅ Tratamento de erros
✅ Performance (< 2s para 95% requisições)
✅ Segurança (OWASP Top 10)
✅ Compatibilidade (browsers, devices)

## 📐 FORMATO DE SAÍDA

Retorne APENAS JSON válido (sem markdown):

{
  "analise_requisitos": {
    "entidades": ["User", "Order", "Product"],
    "regras_negocio": ["Pedido mínimo R$ 50", "Estoque > 0"],
    "riscos": ["Race condition em checkout", "SQL injection em busca"],
    "integrações": ["API Pagamento", "API Correios"]
  },
  
  "cenarios_funcionais": [
    {
      "id": "TC-FUNC-001",
      "titulo": "Criar pedido com produto válido deve gerar ID único",
      "categoria": "CRUD",
      "prioridade": "Crítica",
      "tecnica_teste": "Fluxo Principal",
      "descricao": "Valida criação de pedido com dados válidos",
      "pre_condicoes": [
        "Usuário autenticado com token JWT válido",
        "Produto PROD-001 existe com estoque >= 1"
      ],
      "dados_teste": {
        "usuario": {"id": "USR-001", "email": "test@test.com"},
        "produto": {"id": "PROD-001", "preco": 100.00, "estoque": 5}
      },
      "passos": [
        "1. POST /api/orders com body: {userId: 'USR-001', productId: 'PROD-001', quantity: 1}",
        "2. Verificar response status 201",
        "3. Validar response body contém orderId (formato UUID)",
        "4. Consultar GET /api/orders/{orderId} e validar dados"
      ],
      "resultado_esperado": "Pedido criado com ID único, status 'PENDING', total R$ 100.00",
      "criterios_aceitacao": [
        "Response status = 201",
        "orderId é UUID válido",
        "Estoque do produto decrementado para 4",
        "Tempo de resposta < 500ms"
      ],
      "pos_condicoes": "Pedido existe no banco, estoque atualizado"
    }
  ],
  
  "cenarios_negativos": [...],
  "cenarios_borda": [...],
  "cenarios_seguranca": [...],
  "bugs_provaveis": [...],
  "matriz_rastreabilidade": [...],
  "metricas_qualidade": {...}
}

## ⚠️ REGRAS CRÍTICAS

1. IDs seguem padrão: TC-{TIPO}-{NNN}
2. Prioridades: Crítica|Alta|Média|Baixa
3. Passos são executáveis (não vagos)
4. Resultados são mensuráveis
5. JSON válido (sem trailing commas)

## 📚 EXEMPLO DE QUALIDADE

❌ RUIM: "Testar se login funciona"
✅ BOM: "Login com email válido e senha correta deve retornar token JWT e redirecionar para /dashboard em < 2s"

Retorne APENAS o JSON, sem texto adicional."""

def build_prompt():
    print("🧩 Construindo prompt...")

    with open(SAIDA_JSON, "r", encoding="utf-8") as f:
        requisitos = json.load(f)

    combined = ""
    for section in requisitos["sections"]:
        combined += f"\nSEÇÃO: {section['title']}\n"
        for req in section["requirements"]:
            combined += f"- {req['text']}\n"

    return QA_PROMPT + "\n\nREQUISITOS ANALISADOS:\n" + combined

# ==================================
#  3. LIMPAR JSON VINDO DO GEMINI
# ==================================

def limpar_json_bruto(texto):
    try:
        match = re.search(r'\{.*\}', texto, re.DOTALL)
        return match.group(0) if match else texto
    except:
        return texto

# ====================
#  4. CHAMAR GEMINI
# ====================

def gerar_cenarios(prompt):
    print("🤖 Chamando Gemini 2.0 Flash...")
    resp = client.models.generate_content(
        model="models/gemini-2.0-flash", #caso queira utilizar outro modelo do gemini basta trocar por outro. Ex: "gemini-2.0-flash-lite"
        contents=prompt
    )
    return resp.text

# ===========================================================================================
#  5. FORMATAR CÉLULAS DO WORD (VERSÃO INICIAL DO TEMPLATE DO WORD AINDA PODE SER MELHORADO)
# ===========================================================================================

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")

    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "808080")
        borders.append(el)

    tcPr.append(borders)

def style_header(cell):
    set_cell_bg(cell, "D9D9D9")
    set_cell_borders(cell)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)

# =====================================================================
#  6. GERAR TEMPLATE DO WORD PREENCHIDO COM DADOS ANALISADO PELO GEMINI
# =====================================================================

def salvar_word(json_data):
    print("📝 Gerando documento Word...")

    doc = docx.Document()
    doc.add_heading("Cenários de Teste - QA Automático", level=1)

    def add_table(title, itens):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells

        headers = ["ID", "Título", "Descrição", "Pré-condições", "Passos", "Resultado Esperado"]

        for i, h in enumerate(headers):
            hdr[i].text = h
            style_header(hdr[i])

        for item in itens:
            row = table.add_row().cells
            row[0].text = item.get("id", "")
            row[1].text = item.get("titulo", "")
            row[2].text = item.get("descricao", "")
            row[3].text = "\n".join(item.get("pre_condicoes", []))
            row[4].text = "\n".join(item.get("passos", []))
            row[5].text = item.get("resultado_esperado", "")

            for c in row:
                set_cell_borders(c)

        doc.add_paragraph("")

    if "cenarios_funcionais" in json_data:
        add_table("Cenários Funcionais", json_data["cenarios_funcionais"])

    if "cenarios_negativos" in json_data:
        add_table("Cenários Negativos", json_data["cenarios_negativos"])

    if "cenarios_borda" in json_data:
        add_table("Cenários de Borda", json_data["cenarios_borda"])

    doc.save(WORD_OUTPUT)
    print(f"✅ Documento gerado: {WORD_OUTPUT}")

# ============================================================
#  7. EXECUÇÃO PRINCIPAL
# ============================================================

if __name__ == "__main__":
    extrair_requisitos_docx()

    prompt = build_prompt()
    resposta = gerar_cenarios(prompt)

    resposta_limpa = limpar_json_bruto(resposta)

    with open(TESTES_JSON, "w", encoding="utf-8") as f:
        f.write(resposta_limpa)

    try:
        json_data = json.loads(resposta_limpa)
    except:
        print("❌ Erro: Gemini retornou JSON inválido.")
        print("JSON bruto salvo em testes.json")
        exit()

    salvar_word(json_data)

    print("🎉 Processo concluído com sucesso!")